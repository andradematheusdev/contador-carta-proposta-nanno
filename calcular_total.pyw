import docx
import PySimpleGUI as sg
import ctypes
import re

# Vars Pre Declaration
complete_text = []
values_to_sum = [] 

#Config Vars
identifier = "R$"
ignore = "!R$"

# Main Method
def Start():
    layout = [  
        [sg.Text("Carta Proposta")],
        [sg.Text('Caminho do Arquivo', size=(15, 1)), sg.InputText(key='path'), sg.FileBrowse('Procurar')],
        [sg.Button("Calcular",key="calcular")]
        ]
    w = sg.Window('Calcular Total da Proposta', layout)
    event, values = w.read()
    if event == "calcular" :
        ReadDocument(values['path'])    

def validate(value):
    validation = {}
    exclude = []
    result = 0.0
    for v in value:
        if re.match(r'^-?\d+(?:\.\d+)$', v) is None:
            if "total" in v or "TOTAL" in v:
                pass
            else:
                exclude.append(v)
        else:
            result += float(v)
    validation = {"exclude" : exclude, "result" : result}
    return validation

# Reads the document contents and process it
def ReadDocument(path):
    
    if '\"' in path or "\'" in path:
        path = path.strip('"')
        path = path.strip("'")
    
    # Check if the provided path is empty, null or valid
    if not path == None and not path == "":
        try:
            doc = docx.Document(path)
            tables = doc.tables
            
            # Check if document has a table or only paragraphs  
            if bool(tables):      
                for table in tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                complete_text.append(paragraph.text)
            else:
                for paragraph in doc.paragraphs:
                    complete_text.append(paragraph.text)
                    
        except:            
            ctypes.windll.user32.MessageBoxW(0, "Não foi possível processar o arquivo! Tente novamente.", "Alerta", 0)
            Start()              
    else:
        ctypes.windll.user32.MessageBoxW(0, "O caminho não pode estar vazio!", "Alerta", 0)
        Start()              
            
    # Create the list with the elements
    i = 0
    for index, value in enumerate(complete_text):
        if ignore in complete_text[index]:
            pass
        elif identifier in complete_text[index]:
            i += 1
            if i % 2 == 0:
                global values_to_sum
                values_to_sum.append(complete_text[index])
            
    # Prepare the retrived list string values for float conversion   
    values_to_sum = [s.replace(identifier + " ", "") for s in values_to_sum]
    values_to_sum = [s.replace(".", "") for s in values_to_sum]
    values_to_sum = [s.replace(",", ".") for s in values_to_sum]
            
    # Validate the float data and apply math
    validation = validate(values_to_sum)
    result = validation["result"]
    exclude = validation["exclude"]
    result = "{:,.2f}".format(result)    
    result = str(result)    
    result = result.replace(".", ",")
            
    # Show window result or prints it in console
    if 'doc' in vars():
        ShowResult(result, exclude)

# Creates de result window      
def ShowResult(result, exclude = False):
    if(exclude):
        layout = [
                [sg.Text("Um ou mais valores ou linhas estão com problemas:", size=(30, None))],
                [sg.Text(str(exclude), size=(30, None))],  
                [sg.Button("OK",key="OK")]
                ]
    else:
        layout = [  
                [sg.Text("R$ " + str(result), size=(30, None))], 
                [sg.Button("OK",key="OK")]
                ]
    result_window = sg.Window('Soma Total', layout)
    event, values = result_window.read()
    if event == "OK" or event == result_window.close():
        Start()
 
# Call the main method 
Start()

# Project dev path for references
# G:\Workspace\Python\Nanno Helper